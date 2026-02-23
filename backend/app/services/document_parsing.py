"""Document parsing service - extract text from various file formats."""
import io
import logging
from typing import Optional
import pdfplumber
from pdf2image import convert_from_bytes
import pytesseract
from PIL import Image
from openpyxl import load_workbook
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class DocumentParsingError(Exception):
    """Raised when document parsing fails."""

    pass


class DocumentParsingService:
    """Extract text from PDF, Excel, Word, and image files."""

    # Maximum text to send to AI (token limit consideration)
    MAX_TEXT_LENGTH = 8000

    @staticmethod
    def extract_text(
        file_content: bytes,
        mime_type: str,
    ) -> str:
        """
        Route to appropriate extraction method based on MIME type.

        Args:
            file_content: Raw file bytes
            mime_type: MIME type of file (e.g., 'application/pdf')

        Returns:
            Extracted text content

        Raises:
            DocumentParsingError: If extraction fails or unsupported type
        """
        try:
            if mime_type == "application/pdf":
                return DocumentParsingService.extract_text_from_pdf(file_content)
            elif mime_type in [
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "application/vnd.ms-excel",
            ]:
                return DocumentParsingService.extract_text_from_excel(file_content)
            elif mime_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            ]:
                return DocumentParsingService.extract_text_from_word(file_content)
            elif mime_type.startswith("image/"):
                return DocumentParsingService.extract_text_from_image(file_content)
            else:
                raise DocumentParsingError(f"Unsupported MIME type: {mime_type}")

        except DocumentParsingError:
            raise
        except Exception as e:
            raise DocumentParsingError(f"Failed to extract text: {str(e)}")

    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """
        Extract text from PDF using pdfplumber, with OCR fallback.

        Args:
            file_content: Raw PDF bytes

        Returns:
            Extracted text from all pages

        Raises:
            DocumentParsingError: If extraction fails
        """
        try:
            text_parts = []

            # Try pdfplumber first (faster for text-based PDFs)
            try:
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        text = page.extract_text()
                        if text and text.strip():
                            text_parts.append(f"--- Page {page_num} ---\n{text}")

                extracted = "\n\n".join(text_parts)

                # If we got minimal text, try OCR
                if len(extracted.strip()) < 100:
                    logger.info("PDF has minimal text, attempting OCR fallback")
                    extracted = DocumentParsingService._extract_text_from_pdf_ocr(file_content)

            except Exception as e:
                logger.warning(f"pdfplumber failed ({e}), attempting OCR fallback")
                extracted = DocumentParsingService._extract_text_from_pdf_ocr(file_content)

            # Truncate to max length
            return extracted[: DocumentParsingService.MAX_TEXT_LENGTH]

        except Exception as e:
            raise DocumentParsingError(f"PDF extraction failed: {str(e)}")

    @staticmethod
    def _extract_text_from_pdf_ocr(file_content: bytes) -> str:
        """
        Extract text from PDF using OCR (pytesseract).

        Args:
            file_content: Raw PDF bytes

        Returns:
            Extracted text from all pages

        Raises:
            DocumentParsingError: If OCR fails
        """
        try:
            # Convert PDF to images
            images = convert_from_bytes(file_content)
            text_parts = []

            for page_num, image in enumerate(images, 1):
                text = pytesseract.image_to_string(image)
                if text.strip():
                    text_parts.append(f"--- Page {page_num} ---\n{text}")

            return "\n\n".join(text_parts)

        except Exception as e:
            raise DocumentParsingError(f"PDF OCR extraction failed: {str(e)}")

    @staticmethod
    def extract_text_from_excel(file_content: bytes) -> str:
        """
        Extract text from Excel (.xlsx) files.

        Args:
            file_content: Raw Excel file bytes

        Returns:
            Extracted text from all sheets

        Raises:
            DocumentParsingError: If extraction fails
        """
        try:
            text_parts = []
            workbook = load_workbook(io.BytesIO(file_content))

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"--- Sheet: {sheet_name} ---")

                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    if row_text.strip():
                        text_parts.append(row_text)

            return "\n".join(text_parts)[: DocumentParsingService.MAX_TEXT_LENGTH]

        except Exception as e:
            raise DocumentParsingError(f"Excel extraction failed: {str(e)}")

    @staticmethod
    def extract_text_from_word(file_content: bytes) -> str:
        """
        Extract text from Word (.docx) files.

        Args:
            file_content: Raw Word file bytes

        Returns:
            Extracted text including tables

        Raises:
            DocumentParsingError: If extraction fails
        """
        try:
            text_parts = []
            doc = DocxDocument(io.BytesIO(file_content))

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                text_parts.append("--- Table ---")
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() if cell.text else "" for cell in row.cells
                    )
                    if row_text.strip():
                        text_parts.append(row_text)

            return "\n".join(text_parts)[: DocumentParsingService.MAX_TEXT_LENGTH]

        except Exception as e:
            raise DocumentParsingError(f"Word extraction failed: {str(e)}")

    @staticmethod
    def extract_text_from_image(file_content: bytes) -> str:
        """
        Extract text from image files using OCR.

        Args:
            file_content: Raw image file bytes

        Returns:
            Extracted text using OCR

        Raises:
            DocumentParsingError: If extraction fails
        """
        try:
            image = Image.open(io.BytesIO(file_content))
            text = pytesseract.image_to_string(image)
            return text[: DocumentParsingService.MAX_TEXT_LENGTH]

        except Exception as e:
            raise DocumentParsingError(f"Image OCR extraction failed: {str(e)}")
